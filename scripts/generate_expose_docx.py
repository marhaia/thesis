"""Generate expose_standard.docx from the final exposé content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Resolve paths relative to the repository root so the script is portable.
_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(_REPO_ROOT, "Literature", "notes", "expose", "expose_standard.docx")
IMG = os.path.join(_REPO_ROOT, "Literature", "notes", "pipeline-architecture3.png")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x84, 0x49)

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)

def note(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_after = Pt(4)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = val
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    # Column widths
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()

def blockquote(lines):
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Exposé — Master's Thesis", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = [
    ("Working Title", "Integrating Cognitive Predictive Metrics into the AIM Platform for Automated GUI Evaluation"),
    ("Author", "Hannah"),
    ("Programme", "User Experience Design"),
    ("Supervisors", "Gerhard Graf and Prof. Dr. Christian Sturm"),
    ("Date", "June 2026"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.add_run(label + ": ").bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
h1("1. Introduction and Problem Statement")

doc.add_paragraph(
    "Automated evaluation of graphical user interfaces (GUIs) is a central challenge in HCI research: "
    "designers need scalable, reproducible quality indicators without having to conduct a costly user "
    "study for every design decision. This is particularly critical in safety-relevant contexts such as "
    "automotive HMI, where a poorly designed dashboard can increase cognitive load and compromise driver "
    "safety. The Aalto Interface Metrics (AIM) platform provides an established approach that computes "
    "perception-based metrics such as visual clutter, salience, and symmetry from static screenshots "
    "(Oulasvirta et al., 2018). These metrics are reproducible and objective — but they capture only the "
    "physical properties of the image, not the interplay between interface structure, usage context, and "
    "the user's cognitive capacity. Das et al. (2024) show empirically that cognitive dual-task load "
    "(e.g., simultaneously driving and reading a dashboard) fundamentally alters gaze behaviour: users "
    "exhibit pronounced tunnel vision, miss peripheral UI elements, and form shorter, less efficient "
    "scanpaths. AIM cannot predict this behaviourally relevant phenomenon because it contains neither "
    "task context nor a cognitive model. Jiang et al. (2023) further demonstrate that UI-specific "
    "attention patterns across interface types are systematically underestimated by task-agnostic "
    "saliency models — and propose UMSI++, an improved CNN-based saliency model trained specifically "
    "on diverse UI types, which the present pipeline adopts as its spatial attention module."
)

p = doc.add_paragraph()
p.add_run("The present master's thesis therefore asks: ").italic = False
run = p.add_run(
    "Does the integration of cognitive predictive metrics into the AIM platform improve the "
    "computational evaluation of GUIs?"
)
run.bold = True

# ═══════════════════════════════════════════════════════════════════════════
# 2. RELATED WORK
# ═══════════════════════════════════════════════════════════════════════════
h1("2. Related Work and Theoretical Framework")

h2("2.1 AIM and Computational GUI Evaluation")
doc.add_paragraph(
    "AIM (Oulasvirta et al., 2018) is an open-source platform that computes a range of perception-based "
    "metrics from GUI screenshots. Related tools such as UiLab (Burny & Vanderdonckt, 2021) and "
    "Criticmate (Ko et al., 2026) extend the evaluation framework with consistency checks and multi-stage "
    "critique processes. UIClip (Wu et al., 2024) demonstrates that data-driven CLIP models can learn "
    "GUI quality judgements — but provide no mechanistic explanatory framework and no behavioural prediction. "
    "A systematic review of computational models specifically for automotive HMI design (Lorenz et al., 2024) "
    "finds that none of the 34 surveyed approaches operates on generic screenshots, and only one accounts for "
    "individual user characteristics — precisely the two gaps the present pipeline addresses."
)

h2("2.2 Computational Rationality as Theoretical Foundation")
doc.add_paragraph(
    "The Computational Rationality (CR) framework by Oulasvirta, Jokinen, and Howes (2022) replaces "
    "rule-based engineering models (such as GOMS or ACT-R, which describe interaction as fixed rule "
    "sequences) with sequential decision theory: interaction is modelled as an optimal control policy "
    "that balances cognitive, perceptual, and motor constraints. Formalised as a Partially Observable "
    "Markov Decision Process (POMDP) and solved via Reinforcement Learning, CR explains the transition "
    "from novice-driven, bottom-up search to expert-guided, top-down memory retrieval (Jokinen et al., "
    "2020). Current CR applications include gaze-based selection (Chen et al., 2021), touchscreen typing "
    "(Shi et al., 2024), task-conditioned scanpaths (Shi et al., 2025), attention switching in dual-task "
    "scenarios (Lingler et al., 2024), adaptive UI via multi-agent RL (Langerak et al., 2024), "
    "driver adaptation to assistance systems (Jokinen & Kujala, 2021), and dual-task attention allocation "
    "on head-mounted displays (Bai et al., 2024) — the latter two directly establishing the automotive "
    "dual-task context as a CR-tractable problem. Critically for Stage 2, the Adaptive Feature Guidance "
    "model (Jokinen et al., 2020) formalises visual search as an expected-utility maximisation over fixation "
    "targets, where utility is a function of visual salience, layout homogeneity, and memory retrieval cost. "
    "The ordinal task-context modifiers applied in Stage 2 are grounded in precisely these utility gradients: "
    "high-homogeneity interfaces reduce search cost, while high-salience-dispersion layouts increase it — "
    "a relationship the model makes mathematically explicit."
)

h2("2.3 Saliency, Scanpaths, and Cognitive Load")
doc.add_paragraph(
    "Scanpath prediction has evolved from task-agnostic saliency models (DeepGaze II, Kümmerer et al., "
    "2017) towards task-conditioned approaches: SeekUI (Guo et al., 2026) and Chartist (Shi et al., 2025) "
    "show that task-specific conditioning significantly improves prediction quality. EyeFormer (Jiang et "
    "al., 2024) demonstrates the feasibility of personalised scanpaths through few-shot conditioning. "
    "That cognitive load measurably reshapes scanpaths is empirically established by Krejtz et al. (2018), "
    "who show that fixation-based pupillometry and microsaccade magnitude reliably discriminate task "
    "difficulty (η²≈0.16–0.17) — validating eye-tracking as a cognitive load measure for both the "
    "HCEye dataset and the planned user study. Das et al. (2024) provide the most direct empirical "
    "precedent for the pipeline's target domain: in automotive dual-task conditions, cognitive load "
    "produces tunnel vision, shorter fixations, and missed peripheral elements — the exact behavioural "
    "pattern Stage 2 is designed to estimate. Stage 2 closes this gap mechanistically: UMSI++ produces a "
    "globally-computed saliency map that, by design, has no cognitive capacity limits. The HCEye-calibrated "
    "sensitivity features h ∈ ℝ⁶ act as a load-proportional filter on this global attention distribution — "
    "suppressing peripheral saliency mass in proportion to the predicted cognitive load index. This "
    "approximates the tunnel-vision constriction empirically documented by Das et al. (2024) without "
    "requiring real-time eye-tracking. Cognitive load itself is most commonly operationalised via "
    "the NASA Task Load Index (Hart & Staveland, 1988), a multi-dimensional subjective workload measure "
    "that serves as the validation target for the Cognitive Load Index output of Stage 2."
)

h2("2.4 Research Gap")
doc.add_paragraph(
    "No existing work combines: (a) task-independent complexity extraction as an interpretable input "
    "layer, (b) task-conditioned behavioural prediction with (c) an explicit post-prediction "
    "cross-validation layer to check the behavioural coherence between structural, attentional, and "
    "cognitive load outputs — integrated into an established GUI evaluation platform (AIM)."
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════
h1("3. Research Design and Methodology")
h2("3.1 Architecture: Two-Stage Pipeline")

doc.add_paragraph(
    "The pipeline converts a GUI screenshot into three behavioural estimates. "
    "Figure 1 shows the overall data flow. A screenshot is processed by two parallel modules — "
    "a task-agnostic visual complexity extractor (Stage 1) and a deep saliency model (UMSI++) — "
    "whose outputs are combined into an extended feature vector. This vector is then passed to "
    "Stage 2, which uses task context and an optional user profile to predict three interaction "
    "quality indicators. In the current training-free implementation, Stage 2 predictions are "
    "derived from HCEye-calibrated empirical coefficients; once validation study data are "
    "available, a trained regression model replaces this component. The decoupled architecture "
    "ensures that Stage 1 outputs are reusable across task contexts and that each module can "
    "be updated independently."
)

doc.add_paragraph(
    "The pipeline's three outputs are: a spatial attention map (UMSI++), search efficiency and "
    "attention demand scalars, and a Cognitive Load Index (CLI, 0–100 scale). The CLI is explicitly "
    "not equivalent to a NASA-TLX score: it estimates interactional complexity as experienced through "
    "the interface structure and task context, not subjective workload dimensions such as frustration "
    "or self-assessed performance. Correlation with NASA-TLX serves as the validation target, not as "
    "an identity claim."
)

# Pipeline figure
if os.path.exists(IMG):
    doc.add_picture(IMG, width=Cm(15))
    cap = doc.add_paragraph("Figure 1: Two-Stage Pipeline Architecture")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

h3("Stage 1 — Task-Independent Visual Complexity Extraction")
doc.add_paragraph(
    "Stage 1 computes eight perceptual features from a raw screenshot using established "
    "image-analysis models from the GUI evaluation and visual perception literature; the majority "
    "are implemented within the open-source AIM framework (Oulasvirta et al., 2018). "
    "Output: visual complexity vector v ∈ ℝ⁸."
)

add_table(
    headers=["Feature", "Description", "Source"],
    rows=[
        ["Shannon Entropy",     "Information density (greyscale histogram)",           "Shannon (1948)"],
        ["Edge Density",        "Proportion of edge pixels via Canny",                 "Standard image processing"],
        ["Feature Congestion",  "Multi-channel clutter: colour, contrast, orient.",    "Rosenholtz et al. (2007)"],
        ["Subband Entropy",     "Redundancy via steerable pyramid decomposition",      "Rosenholtz et al. (2007)"],
        ["Layout Symmetry",     "Vertical/horizontal balance of elements",             "Miniukovich & De Angeli (2015)"],
        ["Chromatic Coherence", "Colour palette variance",                             "Hasler & Süsstrunk (2003)"],
        ["Visual Hierarchy",    "Size gradient, contrast, grouping structure",         "Tuch et al. (2009)"],
        ["Element Density",     "Number of interactive controls per area",             "AIM framework"],
    ],
    col_widths=[4.5, 7.5, 4.5],
)

h3("Saliency Module (UMSI++) — Spatial Attention Prediction")
doc.add_paragraph(
    "In parallel, the original screenshot is passed through the UMSI++ deep saliency model "
    "(Jiang et al., 2023; CNN-based), which produces a 2D saliency map. Five scalar features are extracted from "
    "this map (s ∈ ℝ⁵: dispersion, peak count, centre bias, entropy, coverage) and appended to v. "
    "The saliency map is delivered directly by UMSI++ (zero-shot, no retraining); it constitutes "
    "the pipeline's spatial attention output and is validated against eye-tracking fixation "
    "densities to assess how well the pre-trained model transfers to automotive HMI screenshots. "
    "Stage 2 does not retrain or modify UMSI++; it operates on the derived scalar features only."
)

h3("Stage 2 — Task-Conditioned Prediction")
doc.add_paragraph(
    "Stage 2 converts x ∈ ℝ¹⁹ = [v ∈ ℝ⁸ | s ∈ ℝ⁵ | h ∈ ℝ⁶] + Task Descriptor + User Profile "
    "into three interaction quality scores. Note: h is computed directly from the screenshot image "
    "using empirical sensitivity coefficients derived from the HCEye dataset — no participant "
    "eye-tracking data is required at inference time."
)

p = doc.add_paragraph()
p.add_run("Current implementation (training-free): ").bold = True
p.add_run(
    "Predictions are derived directly from the HCEye-calibrated sensitivity features h ∈ ℝ⁶ "
    "via empirically grounded coefficient mappings (Das et al., 2024). This requires no training "
    "data and is fully functional as a screening tool from day one. Task Descriptor and User "
    "Profile apply ordinal modifiers on top of the base prediction."
)

p = doc.add_paragraph()
p.add_run("Planned extension (requires validation study data): ").bold = True
p.add_run(
    "Once empirical ground-truth labels are available from the validation study (Option A or C), "
    "the coefficient-based predictions are replaced by a trained multi-output regression model "
    "(Ridge Regression as baseline, Random Forest and XGBoost as primary estimators, "
    "cross-validated). The model architecture is already implemented and the feature vector "
    "x ∈ ℝ¹⁹ serves as its input unchanged."
)

add_table(
    headers=["Output", "Source", "Validation Metric"],
    rows=[
        ["Saliency Map (2D, spatial attention)",           "UMSI++ — pre-trained, zero-shot",                             "AUC-Judd, NSS, SIM, CC (domain transfer)"],
        ["Search Efficiency + Attention Demand (scalars)", "HCEye rule-based (current) / Stage 2 regression (planned)", "Correlation with eye-tracking AOI metrics"],
        ["Cognitive Load Index (0–100)",                  "HCEye rule-based (current) / Stage 2 regression (planned)", "Correlation with NASA-TLX"],
    ],
    col_widths=[4.5, 6.0, 6.0],
)
doc.add_paragraph(
    "UMSI++ is applied without retraining. Stage 2 regression is activated only after "
    "validation study data are collected."
)

h3("Coherence Check")
doc.add_paragraph(
    "After prediction, a rule-based coherence check flags physically inconsistent output "
    "combinations. Three rules are implemented, each grounded in empirical literature: "
    "(1) high saliency dispersion with implausibly low fixation count (Rosenholtz et al., 2007; "
    "Jokinen et al., 2020); (2) concentrated saliency with high cognitive load (Das et al., "
    "2024; Tuch et al., 2009); (3) high predicted search time with low cognitive load index "
    "(Hart & Staveland, 1988). Incoherent combinations are flagged as warnings; they do not "
    "invalidate the prediction but indicate the input is at the boundary of the model's "
    "calibration range. This rule-based consistency layer is the structural novelty of the "
    "pipeline: no existing GUI evaluation tool cross-validates its outputs across multiple "
    "behavioural dimensions."
)

doc.add_paragraph(
    "User Profile: Big Five trait-level presets (Neutral, Focused/Conscientious, Exploratory/Open, "
    "Social/Extraverted, Stress-Sensitive) modulate the Stage 2 prediction via ordinal weight adjustments. "
    "LLM-based simulation at "
    "trait level is methodologically defensible (Serapio-García et al., 2023; Argyle et al., "
    "2023). State-based simulation (acute stress, fatigue) is explicitly declared as "
    "non-validated (Kapania et al., 2025)."
)

h2("3.2 Empirical Validation")

p = doc.add_paragraph()
p.add_run("Data Source 1 — HCEye Dataset (Das et al., 2024; publicly available):").bold = True
for item in [
    "N=27, eye-tracking + dual-task cognitive load manipulation, 150 web UI screenshots",
    "Use: exploratory calibration of Stage 1 feature norms; benchmark for UMSI++ zero-shot saliency transfer validation",
    "Limitation: web UIs, not automotive HMI — transferability to the target domain is an open question",
]:
    bullet(item)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Data Source 2 — Planned Empirical Validation:").bold = True

blockquote([
    "Open question — supervisor input requested: The form of empirical validation has not yet been "
    "decided. The following options are under consideration; the final choice depends on feasibility, "
    "lab access, and timing, which I would like to discuss.",
    "",
    "Option A — User Study (N≈15–30, within-subjects, eye-tracking + NASA-TLX, automotive HMI "
    "simulator): provides quantitative validation of predictive accuracy and enables the variance "
    "decomposition between Stage 1 and Stage 2.",
    "",
    "Option B — Expert Interviews (N≈5–12 automotive UX/HMI practitioners, structured walkthrough "
    "protocol): assesses plausibility and practical utility of the pipeline outputs from a design perspective.",
    "",
    "Option C — Sequential Mixed Methods (reduced user study + focused expert panel): combines "
    "quantitative prediction validation with qualitative practitioner assessment, at the cost of "
    "greater organisational effort.",
    "",
    "No decision has been made. I would welcome your guidance on which approach is most appropriate "
    "given the timeline and available resources.",
])

p = doc.add_paragraph()
p.add_run("Baseline Comparisons").bold = True
p.add_run(" (require Option A or C — quantitative data collection):")
add_table(
    headers=["Baseline", "Description"],
    rows=[
        ["B1", "Stage 1 alone + linear regression"],
        ["B2", "Stage 2 without coherence check"],
        ["B3", "Single-output model (Cognitive Load Index only)"],
    ],
    col_widths=[3.0, 13.5],
)

p = doc.add_paragraph()
p.add_run("Core Analysis").bold = True
p.add_run(" (require Option A or C):")
bullet("Variance decomposition: how much does Stage 1 explain alone? What does Stage 2 add?")
bullet("Coherence gain: B2 vs. full Stage 2")
bullet("Null results will be reported explicitly")
doc.add_paragraph(
    "If Option B (expert interviews) is chosen, baseline comparisons and quantitative model "
    "evaluation are replaced by a structured practitioner assessment of prediction plausibility "
    "and design utility."
)

# ═══════════════════════════════════════════════════════════════════════════
# 4. CONTRIBUTIONS
# ═══════════════════════════════════════════════════════════════════════════
h1("4. Expected Contributions")
h2("Academic Contributions")

for label, text in [
    ("Theoretical:", "Empirical test of the Computational Rationality hypothesis that cognitive load "
     "can be predicted as the outcome of an optimal control policy over visual interface properties."),
    ("Empirical (contingent on validation Option A or C):", "Variance decomposition of task-independent visual complexity (Stage 1) vs. "
     "task-conditioned behavioural prediction (Stage 2) in a controlled automotive HMI setting — "
     "establishing which layer of the pipeline carries explanatory weight. If Option B (expert interviews) "
     "is chosen instead, this contribution is replaced by a qualitative plausibility assessment of the "
     "pipeline outputs from a practitioner perspective."),
    ("Methodological:", "First integration of a multi-dimensional GUI evaluation pipeline that "
     "cross-validates structural, attentional, and cognitive load estimates through a "
     "post-prediction, rule-based coherence layer — making divergent model outputs visible "
     "as design warnings rather than silently absorbing them into a joint loss."),
]:
    p = doc.add_paragraph()
    p.add_run(label + " ").bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)

h2("Practical Deliverables")
for label, text in [
    ("Two-Stage Pipeline:", "The pipeline itself is the primary deliverable — functioning as a "
     "pre-deployment screening tool for automotive HMI designers. It provides a scalable first "
     "estimate of cognitive load risk at the layout stage, without requiring eye-tracking "
     "infrastructure or a user study."),
    ("Open-Source AIM Extension:", "The implementation is released as an AIM-compatible module, "
     "enabling adoption as a web service or design tool integration."),
]:
    p = doc.add_paragraph()
    p.add_run(label + " ").bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════
# 5. TIMELINE
# ═══════════════════════════════════════════════════════════════════════════
h1("5. Timeline")
note("Implementation will make use of AI-assisted development (large language model coding tools) "
     "throughout, in line with current practice in computational HCI research.")

add_table(
    headers=["Phase", "Content", "Period"],
    rows=[
        ["Exposé",                        "Finalise and submit exposé; supervisor meeting",                                           "June 2026"],
        ["Literature",                      "Finalize reading, close remaining gaps",                                                   "June 2026"],
        ["Registration",                    "Official thesis registration",                                                            "July 2026"],
        ["Implementation (finalization)",   "Complete remaining 10%, integration testing, AIM compatibility",                          "July 2026"],
        ["Study (Preparation)",             "Ethics approval (parallel), recruitment, stimuli, pilot test",                           "July–Aug 2026"],
        ["Study (Data Collection)",         "Main study N≈15–30 (form TBD, see Section 3.2)",                                          "Aug–Sep 2026"],
        ["Analysis & Training",             "Eye-tracking processing, model training, baseline comparisons B1\u2013B3 (if Option A or C; replaced by expert interview analysis if Option B)",     "Sep\u2013Oct 2026"],
        ["Writing",                         "Thesis chapters 1–7 (begins in parallel with analysis)",                                 "Oct–Nov 2026"],
        ["Submission",                      "—",                                                                                       "End of Dec 2026"],
    ],
    col_widths=[4.5, 9.0, 3.0],
)

# ═══════════════════════════════════════════════════════════════════════════
# 6. LIMITATIONS
# ═══════════════════════════════════════════════════════════════════════════
h1("6. Scope and Limitations")
for item in [
    "Domain restriction: automotive HMI screenshots only; generalisation to other UI types requires separate validation",
    "Cognitive Load Index ≠ NASA-TLX; the output is a computed interaction complexity indicator",
    "LLM user simulation: trait-level only (Big Five); state-based simulation (fatigue, acute stress) is not validated "
    "and should not be interpreted as such — the User Profile feature is exploratory; core validation rests on the "
    "structural and task-based inputs to Stage 2 (Kapania et al., 2025)",
    "Pilot dataset: single-GUI design without complexity variation; suitable for calibration only",
]:
    bullet(item)

# ═══════════════════════════════════════════════════════════════════════════
# 7. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════
h1("7. References")

# APA-style: Author(s). (Year). Title. Venue. https://doi.org/...
refs = [
    "Argyle, L. P., Busby, E. C., Fulda, N., Gubler, J. R., Rytting, C., & Wingate, D. (2023). Out of one, many: Using language models to simulate human samples. Political Analysis. https://doi.org/10.1017/pan.2023.2",
    "Bai, Y., Ikkala, A., Oulasvirta, A., Zhao, S., Wang, L. J., Yang, P., & Xu, P. (2024). Heads-up multitasker: Simulating attention switching on optical head-mounted displays. Proceedings of CHI 2024. https://doi.org/10.1145/3613904.3642540",
    "Burny, N., & Vanderdonckt, J. (2021). UiLab, a workbench for conducting and replicating experiments in GUI visual design. Proceedings of EICS 2021. https://doi.org/10.1145/3457143",
    "Chen, X., Acharya, A., Oulasvirta, A., & Howes, A. (2021). An adaptive model of gaze-based selection. Proceedings of CHI 2021. https://doi.org/10.1145/3411764.3445177",
    "Das, A., Wu, Z., Škrjanec, I., & Feit, A. M. (2024). Shifting focus with HCEye: Exploring the dynamics of visual highlighting and cognitive load on user attention and saliency prediction. Proceedings of ETRA 2024. https://doi.org/10.1145/3655610",
    "Guo, Z., Jiang, Y., Leiva, L. A., & Oulasvirta, A. (2026). SeekUI: Predicting visual search behavior on graphical user interfaces with a reward-augmented vision language model. Proceedings of CHI 2026. https://doi.org/10.1145/3772318.3791178",
    "Hart, S. G., & Staveland, L. E. (1988). Development of NASA-TLX (Task Load Index): Results of empirical and theoretical research. In P. A. Hancock & N. Meshkati (Eds.), Human Mental Workload. https://doi.org/10.1016/S0166-4115(08)62386-9",
    "Hasler, D., & Süsstrunk, S. E. (2003). Measuring colourfulness in natural images. SPIE Human Vision and Electronic Imaging VIII. https://doi.org/10.1117/12.477378",
    "Jiang, Y., Leiva, L. A., Houssel, P. R. B., Tavakoli, H. R., Kylmälä, J., & Oulasvirta, A. (2023). UEyes: Understanding visual saliency across user interface types. Proceedings of CHI 2023. https://doi.org/10.1145/3544548.3581113",
    "Jiang, Y., Guo, Z., Tavakoli, H. R., Leiva, L. A., & Oulasvirta, A. (2024). EyeFormer: Predicting personalized scanpaths with transformer-guided reinforcement learning. Proceedings of UIST 2024. https://doi.org/10.1145/3654777.3676436",
    "Jokinen, J. P. P., Wang, Z., Sarcar, S., Oulasvirta, A., & Ren, X. (2020). Adaptive feature guidance: Modelling visual search with graphical layouts. International Journal of Human-Computer Studies, 136, 102376. https://doi.org/10.1016/j.ijhcs.2019.102376",
    "Jokinen, J. P. P., & Kujala, T. (2021). Modelling drivers' adaptation to assistance systems. Proceedings of AutomotiveUI 2021. https://doi.org/10.1145/3409118.3475150",
    "Kapania, S., Agnew, W., Eslami, M., Heidari, H., & Fox, S. E. (2025). Simulacrum of stories: Examining large language models as qualitative research participants. Proceedings of CHI 2025. https://doi.org/10.1145/3706598.3713220",
    "Ko, J., Choi, J., Morales, C., Kim, D., & Ko, M. (2026). Criticmate: Stagewise human–AI co-critique in single-screen UI evaluation. Proceedings of CHI 2026. https://doi.org/10.1145/3772318.3790929",
    "Krejtz, K., Duchowski, A. T., Niedzielska, A., Biele, C., & Krejtz, I. (2018). Eye tracking cognitive load using pupil diameter and microsaccades with fixed gaze. PLOS ONE, 13(9), e0203629. https://doi.org/10.1371/journal.pone.0203629",
    "Kümmerer, M., Wallis, T. S. A., Gatys, L. A., & Bethge, M. (2017). Understanding low- and high-level contributions to fixation prediction. ICCV 2017. https://doi.org/10.1109/ICCV.2017.513",
    "Langerak, T., Christen, S., Albaba, M., Gebhardt, C., Holz, C., & Hilliges, O. (2024). MARLUI: Multi-agent reinforcement learning for adaptive point-and-click UIs. ACM Transactions on Computer-Human Interaction. https://doi.org/10.1145/3661147",
    "Lingler, A., Talypova, D., Jokinen, J. P. P., Oulasvirta, A., & Wintersberger, P. (2024). Supporting task switching with reinforcement learning. Proceedings of CHI 2024. https://doi.org/10.1145/3613904.3642063",
    "Lorenz, M., Amorim, T., Dey, D., Sadeghi, M., & Ebel, P. (2024). Computational models for in-vehicle user interface design: A systematic literature review. Proceedings of AutomotiveUI 2024. https://doi.org/10.1145/3640792.3675735",
    "Miniukovich, A., & De Angeli, A. (2015). Computation of interface aesthetics. Proceedings of CHI 2015. https://doi.org/10.1145/2702123.2702575",
    "Oulasvirta, A., De Pascale, S., Koch, J., Langerak, T., Jokinen, J., Todi, K., Laine, M., Kristhombuge, M., Zhu, Y., Miniukovich, A., Palmas, G., & Weinkauf, T. (2018). Aalto Interface Metrics (AIM): A service and codebase for computational GUI evaluation. Adjunct Proceedings of UIST 2018, 16–19. https://doi.org/10.1145/3266037.3266087",
    "Oulasvirta, A., Jokinen, J. P. P., & Howes, A. (2022). Computational rationality as a theory of interaction. Proceedings of CHI 2022. https://doi.org/10.1145/3491102.3517739",
    "Rosenholtz, R., Li, Y., & Nakano, L. (2007). Measuring visual clutter. Journal of Vision, 7(2), 17. https://doi.org/10.1167/7.2.17",
    "Serapio-García, G., Safdari, M., Crepy, C., Sun, L., Fitz, S., Romero, P., Abdulhai, M., Faust, A., & Matarić, M. (2023). Personality traits in large language models. arXiv. https://doi.org/10.48550/arXiv.2307.00184",
    "Shannon, C. E. (1948). A mathematical theory of communication. Bell System Technical Journal, 27(3), 379–423. https://doi.org/10.1002/j.1538-7305.1948.tb01338.x",
    "Shi, D., Zhu, Y., Jokinen, J. P. P., Acharya, A., Putkonen, A., Zhai, S., & Oulasvirta, A. (2024). CRTypist: Simulating touchscreen typing behavior via computational rationality. Proceedings of CHI 2024. https://doi.org/10.1145/3613904.3642918",
    "Shi, D., Wang, Y., Bai, Y., Bulling, A., & Oulasvirta, A. (2025). Chartist: Task-driven eye movement control for chart reading. Proceedings of CHI 2025. https://doi.org/10.1145/3706598.3713128",
    "Tuch, A. N., Bargas-Avila, J. A., Opwis, K., & Wilhelm, F. H. (2009). Visual complexity of websites: Effects on users' experience, physiology, performance, and memory. International Journal of Human-Computer Studies, 67(9), 703–715. https://doi.org/10.1016/j.ijhcs.2009.04.002",
    "Wu, J., Peng, Y.-H., Li, X. Y., Swearngin, A., Bigham, J. P., & Nichols, J. (2024). UIClip: A data-driven model for assessing user interface design. Proceedings of UIST 2024. https://doi.org/10.1145/3654777.3676408",
]

for ref in refs:
    bullet(ref)

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
